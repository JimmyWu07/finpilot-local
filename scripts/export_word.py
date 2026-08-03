# requires: python-docx
# 标准化投研报告Word导出引擎 - Markdown转Word文档
import argparse
import json
import sys
import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print(json.dumps({"error": "python-docx未安装，请执行: pip install python-docx"}))
    sys.exit(1)


def create_cover_page(doc, title, subtitle, date_str):
    """创建封面页"""
    # 空行占位
    for _ in range(6):
        doc.add_paragraph("")
    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    # 日期
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成日期: {date_str}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FinPilot AI 投研助手")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()


def add_header_footer(doc, stock_code, company_name, date_str):
    """添加页眉页脚"""
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = f"FinPilot 投研报告 | {stock_code} {company_name} | {date_str}"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = "仅供研究参考，不构成投资建议"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)


def parse_markdown_table(lines, start_idx):
    """解析Markdown表格"""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells and not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def markdown_to_docx(md_content, doc, stock_code="", company_name=""):
    """将Markdown内容转换为Word文档"""
    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # 空行
        if not stripped:
            i += 1
            continue
        # 标题
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped.lstrip("#").strip()
            # 移除emoji
            title_text = re.sub(r'[📊📈📉⚠️💡🔵🟡🟠🔴🟢]', '', title_text).strip()
            heading = doc.add_heading(title_text, level=min(level, 4))
            i += 1
            continue
        # 表格
        if stripped.startswith("|"):
            rows, end_idx = parse_markdown_table(lines, i)
            if rows:
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < max_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                doc.add_paragraph("")
            i = end_idx
            continue
        # 列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = re.sub(r'[📊📈📉⚠️💡🔵🟡🟠🔴🟢]', '', text).strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        # 普通段落
        text = re.sub(r'[📊📈📉⚠️💡🔵🟡🟠🔴🟢]', '', stripped).strip()
        if text:
            # 处理加粗
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        i += 1


def export_report(md_path, output_path, stock_code="", company_name=""):
    """导出报告"""
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    date_str = datetime.now().strftime("%Y-%m-%d")
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    # 封面
    title = "投资研究分析报告"
    subtitle = f"{company_name or '目标公司'} ({stock_code or '股票代码'})"
    create_cover_page(doc, title, subtitle, date_str)
    # 页眉页脚
    add_header_footer(doc, stock_code, company_name, date_str)
    # 转换内容
    markdown_to_docx(md_content, doc, stock_code, company_name)
    # 保存
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    return {
        "status": "success",
        "output_path": output_path,
        "file_size": file_size,
        "file_size_human": f"{file_size/1024:.1f}KB"
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="投研报告Word导出")
    parser.add_argument("--input", required=True, help="输入Markdown文件路径")
    parser.add_argument("--output", required=True, help="输出Word文件路径")
    parser.add_argument("--stock_code", default="", help="股票代码")
    parser.add_argument("--company_name", default="", help="公司名称")
    args = parser.parse_args()
    result = export_report(args.input, args.output, args.stock_code, args.company_name)
    print(json.dumps(result, ensure_ascii=False))
